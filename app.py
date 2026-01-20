import streamlit as st
from PIL import Image
from docx import Document
from io import BytesIO
from docx.shared import Pt, Cm

# ================== CẤU HÌNH ==================
st.set_page_config(
    page_title="Trợ lý GV tiểu học – Lời dẫn lên lớp",
    page_icon="📘",
    layout="centered"
)

# ================== TIÊU ĐỀ ==================
st.markdown("""
<div style="text-align:center;">
    <h1>📘 TRỢ LÝ GIÁO VIÊN TIỂU HỌC</h1>
    <h3>KỊCH BẢN LÊN LỚP – LỜI DẪN CHI TIẾT TỪNG BƯỚC</h3>
    <p><i>Không AI – Không lỗi – Dùng được ngay khi lên lớp</i></p>
    <p style="color:#555;"><b>✍️ Tác giả:</b> NGUYỄN VĂN DU – Giáo viên Tiểu học</p>
</div>
<hr>
""", unsafe_allow_html=True)

# ================== THÔNG TIN ==================
st.markdown("## 📝 THÔNG TIN BÀI DẠY")
mon = st.selectbox("📚 Môn học", ["Tin học", "Công nghệ", "Toán", "Tiếng Việt"])
lop = st.selectbox("🎓 Lớp", ["3", "4", "5"])
ten_bai = st.text_input("📖 Tên bài học")

# ================== ẢNH SGK ==================
st.markdown("## 📸 ẢNH SÁCH GIÁO KHOA (THAM KHẢO)")
uploaded_images = st.file_uploader(
    "Tải hoặc chụp nhiều trang SGK",
    type=["jpg", "jpeg", "png"],
    accept_multiple_files=True
)

if uploaded_images:
    cols = st.columns(3)
    for i, f in enumerate(uploaded_images):
        with cols[i % 3]:
            st.image(Image.open(f), use_column_width=True)

# ================== GV NHẬP NỘI DUNG ==================
st.markdown("## ✍️ GIÁO VIÊN GHI NỘI DUNG TRỌNG TÂM (THEO SGK)")
noidung = st.text_area(
    "Mỗi ý 1 dòng (ghi đúng SGK):",
    height=180,
    placeholder="Ví dụ:\n- Khái niệm máy tính\n- Các bộ phận chính\n- Lợi ích của máy tính"
)

# ================== TẠO KỊCH BẢN ==================
if st.button("🚀 TẠO KỊCH BẢN LÊN LỚP (GV NÓI CHI TIẾT)"):
    if not ten_bai or not noidung:
        st.warning("⚠️ Cần nhập TÊN BÀI và NỘI DUNG")
        st.stop()

    content = f"""
BÀI: {ten_bai}
MÔN: {mon} – LỚP: {lop}

=================================================
I. KHỞI ĐỘNG (5 phút)

🎤 GV nói:
- Các em ổn định chỗ ngồi, chuẩn bị sách vở.
- Trước khi vào bài mới, cô/trò ta cùng trao đổi một chút nhé.
- (GV đặt câu hỏi gợi mở liên quan bài học).

👧👦 HS:
- HS suy nghĩ và trả lời theo hiểu biết cá nhân.

🔁 Nếu HS trả lời chưa đúng:
- GV gợi ý nhẹ nhàng, đặt câu hỏi phụ.

✅ GV chốt:
- Nhận xét câu trả lời của HS.
- Giới thiệu: “Hôm nay chúng ta sẽ học bài: {ten_bai}”.

=================================================
II. HÌNH THÀNH KIẾN THỨC (15 phút)

🎤 GV nói:
- Các em mở sách, quan sát nội dung bài học.
- Cô sẽ hướng dẫn từng phần, các em chú ý lắng nghe.

👉 Nội dung chính:
{noidung}

🎤 GV hỏi:
- Theo các em, nội dung trên cho ta biết điều gì?
- Ai có thể nhắc lại ý chính?

👧👦 HS:
- HS quan sát, suy nghĩ, trả lời.
- Một số HS khác nhận xét, bổ sung.

🔁 Nếu HS lúng túng:
- GV đọc lại ý trong SGK, giải thích bằng lời dễ hiểu.

✅ GV chốt:
- Khẳng định kiến thức đúng.
- Nhấn mạnh nội dung cần ghi nhớ.

=================================================
III. LUYỆN TẬP (10 phút)

🎤 GV nói:
- Bây giờ chúng ta cùng luyện tập để hiểu bài hơn.
- GV nêu câu hỏi/bài tập liên quan nội dung vừa học.

👧👦 HS:
- HS làm việc cá nhân hoặc theo nhóm.
- Trình bày kết quả trước lớp.

🔁 Nếu HS làm sai:
- GV hướng dẫn lại từng bước.
- Cho HS làm lại.

✅ GV chốt:
- Nhận xét chung.
- Tuyên dương HS làm tốt.

=================================================
IV. VẬN DỤNG (5 phút)

🎤 GV nói:
- Các em hãy liên hệ kiến thức vừa học với thực tế.
- GV đặt câu hỏi: “Trong cuộc sống, em đã gặp nội dung này ở đâu?”

👧👦 HS:
- HS nêu ví dụ thực tế.

✅ GV chốt:
- Củng cố lại toàn bài.
- Dặn dò HS về nhà ôn bài, chuẩn bị bài sau.
"""

    st.markdown("## 📄 KỊCH BẢN LÊN LỚP (GV NÓI)")
    st.text(content)

    # ================== XUẤT WORD ==================
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(3)
    section.right_margin = Cm(2)

    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(14)

    doc.add_paragraph("Tác giả: NGUYỄN VĂN DU – Giáo viên Tiểu học").italic = True

    for line in content.split("\n"):
        doc.add_paragraph(line)

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)

    st.download_button(
        "⬇️ Tải file Word – Kịch bản GV nói",
        buf,
        file_name=f"Kich_ban_GV_noi_{ten_bai}.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )

st.markdown("---")
st.markdown(
    "<div style='text-align:center; color:#666;'>© 2026 – Trợ lý giáo viên | Nguyễn Văn Du</div>",
    unsafe_allow_html=True
)